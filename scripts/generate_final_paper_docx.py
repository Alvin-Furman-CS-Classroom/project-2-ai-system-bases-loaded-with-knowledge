#!/usr/bin/env python3
"""
Build an editable CSC343 Project 2 final paper as a .docx (Word) file.

Requires: python-docx (repo-local venv example):
  python3 -m venv .venv_paper && .venv_paper/bin/pip install python-docx
  .venv_paper/bin/python scripts/generate_final_paper_docx.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "CSC343_Project2_Final_Paper.docx"


def _set_body_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.space_after = Pt(6)


def _add_para(doc: Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def main() -> None:
    doc = Document()
    _set_body_style(doc)

    # --- Title block ---
    t = doc.add_paragraph()
    r = t.add_run(
        "AI-Supported Baseball and Softball Lineup Optimization Using Logic, "
        "Knowledge Bases, Constraints, Evolutionary Search, and Planning"
    )
    r.bold = True
    r.font.size = Pt(14)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    _add_para(doc, "Team members: Kali Javetski; Sylvia Burroughs", bold=True)
    _add_para(doc, "Repository URL: [paste your team repository link]", bold=False)
    _add_para(doc, "Demo link (optional): [paste dashboard or demo URL]", bold=False)
    doc.add_paragraph()
    _add_para(
        doc,
        "Note: Before PDF submission, insert page numbers (Insert → Page Number), "
        "replace bracketed placeholders, add Figure 1 architecture art and a "
        "Figure 2 screenshot, run a word count (main text 2200–2500 words per "
        "course spec, abstract excluded), and verify IEEE citations.",
        bold=False,
    )

    doc.add_page_break()

    # --- Abstract ---
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "Constructing a competitive baseball or softball lineup requires reconciling "
        "batter–pitcher matchup expectations, heterogeneous defensive value by position, "
        "hard eligibility constraints, combinatorial batting-order choices, and "
        "shifting in-game context. This project implements a five-stage decision-support "
        "pipeline in standard-library Python: (1) quantified rule-based offensive "
        "projections from season-style batter and pitcher profiles; (2) knowledge-base-"
        "driven defensive grading with optional cross-position inference; (3) constraint "
        "satisfaction with branch-and-bound optimization to assign nine roles including "
        "pitcher; (4) a genetic algorithm that searches permutations under a baseball-"
        "shaped fitness function; and (5) a planning layer that ranks explainable tactical "
        "recommendations over a short lookahead horizon. The system is modular, test-"
        "driven, and integrated end-to-end from statistics files through a browser-"
        "oriented demo path for later modules. Automated evidence consists of 243 passing "
        "unit tests across modules and three integration tests that exercise the full "
        "pipeline on representative fixtures. The primary outcome is not a claim of real-"
        "world win probability improvement—which would require field validation—but a "
        "coherent, inspectable toolchain that maps course topics to a single domain "
        "problem while exposing clear limitations in modeling scope (for example, no "
        "dedicated pitching or bullpen module)."
    )

    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "Youth, amateur, and even professional coaching staffs routinely assemble lineups "
        "under incomplete information: limited samples, uncertain defensive utility at "
        "secondary positions, and rapidly changing score states. Computational assistants "
        "can help if they remain interpretable, respect roster constraints, and separate "
        "pre-game optimization from in-game adaptation."
    )
    doc.add_paragraph(
        "This report describes the system built for CSC343 Project 2: an assistive lineup "
        "tool. It does ingest tabular statistics, produce numeric matchup and defensive "
        "scores, solve a nine-slot assignment, optimize a batting order, and emit "
        "prioritized tactical suggestions tied to game state. It does not simulate "
        "pitch-by-pitch outcomes, learn from proprietary play-by-play corpora, or output "
        "calibrated win-expectancy estimates. Pitcher fatigue fields are accepted for "
        "interface compatibility but are not used by the planner because pitching "
        "strategy is out of scope for the implemented modules, as documented in the "
        "repository README."
    )
    doc.add_paragraph(
        "The motivation is pedagogical and practical: the same domain threads together "
        "first-order logic style rules, structured knowledge evaluation, CSP search, "
        "evolutionary permutation search, and lightweight planning—each technique "
        "occupying a clean boundary with explicit inputs and outputs."
    )
    doc.add_paragraph(
        "From a software engineering perspective, the project also demonstrates how to "
        "keep each AI technique replaceable. Module boundaries are enforced by typed "
        "dictionaries and small public entrypoints (for example, analyze_matchup_performance, "
        "analyze_defensive_performance, assign_defensive_positions, optimize_batting_order, "
        "and generate_adaptive_plan), which makes it feasible to swap a heuristic for a "
        "learned model later without rewriting the entire pipeline."
    )
    doc.add_paragraph(
        "Scope boundaries matter for honest reporting. The system is not a full "
        "simulator: it does not model pitch selection, catcher framing, weather, park "
        "factors, or opponent bullpen sequencing. Those omissions are intentional for a "
        "course-scale artifact, but they also define where recommendations should be "
        "treated as decision support rather than autonomous policy."
    )

    doc.add_heading("2. System Architecture", level=1)
    doc.add_paragraph(
        "Figure 1 shows the end-to-end architecture (replace the placeholder box in the "
        "final PDF with your team's diagram). Data flows forward on the happy path: "
        "statistics files feed Modules 1 and 2 in parallel; their dictionaries merge into "
        "Module 3; Module 3's nine names and stat slices feed Module 4; Module 4's "
        "batting order plus Module 3's field map feed Module 5 together with bench "
        "metadata and flattened score maps required by the planner API."
    )
    doc.add_paragraph(
        "Figure 1. High-level architecture of the five-module pipeline: logic and "
        "knowledge bases → CSP assignment → genetic batting order → planning. "
        "[Replace with final diagram; caption below figure per course instructions.]"
    )
    doc.add_paragraph(
        "The repository layout separates src/module1 through module5, mirrored tests "
        "under unit_tests/, and staged integration under integration_tests/. A static HTML "
        "dashboard (web/module4_dashboard.html, generated via demo scripts in the README) "
        "provides a human-facing view of later pipeline segments without introducing "
        "third-party web frameworks, preserving the standard-library-only dependency "
        "posture described in the README [1]."
    )
    doc.add_paragraph(
        "Architecturally, Modules 1 and 2 are intentionally parallel preprocessors. "
        "They read different columns from roster files and produce different score shapes: "
        "Module 1 emits a single offensive scalar per player for the current opponent "
        "pitcher context, while Module 2 emits a matrix of defensive values keyed by "
        "position. Module 3 is the first stage that requires both signals simultaneously, "
        "because its objective couples per-position defense with a player-level offensive "
        "contribution."
    )
    doc.add_paragraph(
        "Module 4 consumes only the nine-player subset chosen by Module 3 plus a compact "
        "batter stat slice needed for fitness evaluation. Module 5 consumes the richest "
        "state: it needs the batting order and defensive map from upstream modules, but "
        "also requires an explicit game_state object and bench metadata so recommendations "
        "respect substitution limits and inning context. This layering keeps expensive "
        "search (CSP and GA) separated from rapidly iterated tactical rules."
    )

    doc.add_heading("3. Module Implementation Summary", level=1)

    doc.add_paragraph(
        "This section summarizes what shipped in each module, emphasizing interfaces and "
        "the rationale for coupling decisions. The goal is not to restate every function "
        "name, but to show how each course topic manifests as an implementable component "
        "with measurable tests."
    )

    doc.add_heading("3.1 Module 1 — Matchup analysis (first-order logic)", level=2)
    doc.add_paragraph(
        "Purpose: map each batter to a 0–100 offensive score against a fixed opponent "
        "pitcher profile using rule evaluation inspired by quantified matchup heuristics "
        "(handedness, rate stats, and thresholds). Inputs and outputs: JSON or CSV roster "
        "files with batter and pitcher fields; output is a map from batter name to score "
        "consumed by Modules 3–5. Design choices: a dedicated RuleEvaluator composes "
        "adjustments on top of a base calculator, keeping rule logic testable in isolation."
    )
    doc.add_paragraph(
        "Module 1 is deliberately not a machine-learning predictor: it encodes explicit "
        "domain regularities that coaches already verbalize (platoon disadvantage, "
        "discipline versus wildness, power versus mistake-prone pitching). That choice "
        "sacrifices predictive accuracy relative to large-scale learned models, but it "
        "improves inspectability and makes failures easier to diagnose when a score looks "
        "wrong for a known player."
    )

    doc.add_heading("3.2 Module 2 — Defensive analysis (knowledge bases)", level=2)
    doc.add_paragraph(
        "Purpose: produce position-conditioned defensive scores for each player-position "
        "pair, with catcher-specific metrics. Inputs and outputs: defensive statistics "
        "files; output maps each player to scores per position for Module 3. Design "
        "choices: beyond direct observation of played positions, the implementation can "
        "predict performance at unplayed positions using cross-position similarity "
        "heuristics, configurable via predict_all_positions, as documented in the README."
    )
    doc.add_paragraph(
        "Knowledge-base framing matters because defensive value is partly positional: a "
        "shortstop error is not interchangeable with a first baseman error in how it "
        "stresses range and arm strength. Encoding those distinctions as rules keeps the "
        "evaluation aligned with coaching intuition while still producing numeric scores "
        "that CSP can optimize."
    )

    doc.add_heading("3.3 Module 3 — Position assignment (CSP)", level=2)
    doc.add_paragraph(
        "Purpose: assign nine unique players to nine positions maximizing a weighted sum "
        "of defense and offense subject to eligibility and all-different constraints. "
        "Inputs: offensive dict, nested defensive dict, eligibility lists; output: "
        "position-to-player map. Design choices: backtracking with MRV, forward checking, "
        "and LCV-style ordering. An optional defensive stress profile scales defensive "
        "contributions using FanGraphs positional adjustment runs mapped to near-unity "
        "multipliers; the pitcher's defensive term stays neutral because Module 2 does "
        "not grade pitching defense [3]."
    )
    doc.add_paragraph(
        "The CSP wrapper also supports operational constraints that appear in real dugouts, "
        "such as locking a player to a position when a coach has already committed to a "
        "starter at catcher or shortstop. Locks shrink the search space and make the solver "
        "behave more like a decision aid under partial human commitments rather than a fully "
        "autonomous selector."
    )

    doc.add_heading("3.4 Module 4 — Batting order (genetic algorithm)", level=2)
    doc.add_paragraph(
        "Purpose: search the permutation space for a strong batting order under a "
        "baseball-shaped fitness model. Inputs: nine selected names and per-player OBP, "
        "SLG, HR, and RBI; optional blending with Module 1 scores. Outputs include "
        "optimized_order, best_fitness, generations_run, and seed for reproducibility. "
        "Default hyperparameters (population_size=120, generations=250, mutation_rate=0.08, "
        "elite_count=6, stagnation-based early stopping) balance runtime and stability."
    )
    doc.add_paragraph(
        "The genetic representation is a permutation chromosome, and operators preserve "
        "validity (nine unique players). Fitness encodes conventional lineup structure: "
        "reward high on-base skills near the top, reward slugging and run production in "
        "the middle, and avoid obviously mismatched slotting at the bottom. This is a "
        "transparent compromise between sabermetric nuance and assignment-scale "
        "complexity."
    )

    doc.add_heading("3.5 Module 5 — Adaptive planning", level=2)
    doc.add_paragraph(
        "Purpose: translate validated game state, current lineup, bench roles, and score "
        "context into ranked recommendations and a multi-inning outlook. Inputs: "
        "structured game_state keys, current_lineup with batting order and field map, bench "
        "player objects, and offensive and defensive score maps. Design choices: rule-"
        "first scoring in strategy_rules.py with orchestration and deterministic tie "
        "handling in planner.py to keep recommendations explainable."
    )
    doc.add_paragraph(
        "Planning outputs are intentionally verbose for transparency: recommendations carry "
        "human-readable rationales and priority metadata so a user can trace why a pinch "
        "hitter or defensive replacement surfaced. This design favors auditability over "
        "minimal token output, which aligns with classroom demonstration goals and makes "
        "regression tests easier to write than for opaque numeric policies."
    )

    doc.add_heading("4. Evaluation Methodology", level=1)
    doc.add_paragraph(
        "Evaluation emphasizes regression testing and integration contracts rather than "
        "retrospective league forecasting, which would be a separate study with its own "
        "data governance. What was evaluated: (i) unit correctness of each module's public "
        "APIs; (ii) cross-module compatibility on shared fixtures; (iii) planner output "
        "schema on integrated inputs. Metrics: automated test pass rate and per-module "
        "test counts as a proxy for specification coverage. Qualitative rubric reviews are "
        "recorded in checkpoint elegance reports referenced by the README checkpoint table."
    )
    doc.add_paragraph(
        "We also treat developer ergonomics as secondary evidence: the README documents "
        "how to run each module's tests with consistent PYTHONPATH conventions, which "
        "reduces the chance that a reviewer runs tests incorrectly and misattributes a "
        "failure to the code rather than the environment."
    )
    doc.add_paragraph(
        "Test setup: Python 3.9 with PYTHONPATH=src; python3 -m unittest discover run per "
        "module directory, plus explicit integration modules as documented in the README "
        "[1], [2]. Data: curated JSON and CSV fixtures under test_data/ and toy instances "
        "embedded in tests."
    )
    doc.add_paragraph(
        "Negative testing is part of the methodology, not an afterthought. Several "
        "modules include tests for malformed inputs, infeasible eligibility patterns, and "
        "determinism under fixed random seeds for the genetic optimizer. These tests do "
        "not prove competitive superiority on the field, but they do provide repeatable "
        "evidence that the pipeline contracts are stable as the codebase evolves."
    )
    doc.add_paragraph(
        "Integration tests are intentionally few but strict: they validate that real "
        "outputs from earlier stages can be passed into later stages without manual "
        "renaming or ad hoc reshaping. That design choice trades breadth (many scenarios) "
        "for depth (true cross-module wiring), which matches a team-sized project timeline "
        "while still satisfying the course expectation for end-to-end evidence."
    )

    doc.add_heading("5. Results", level=1)
    doc.add_paragraph(
        "Strengths are reproducible at the engineering level. Table 1 summarizes automated "
        "testing evidence collected while preparing this report. Weaknesses are explicit "
        "scope boundaries: heuristic scores are not calibrated run-value units; the genetic "
        "algorithm returns a strong local optimum under its fitness, not a proven global "
        "optimum; the planner does not recommend pitching changes from fatigue because that "
        "signal is unused in the shipped planner."
    )
    doc.add_paragraph(
        "Another strength is runtime footprint: because the implementation avoids heavy "
        "dependencies, the full unittest suite completes quickly on a laptop, which makes "
        "it realistic for a team to run tests frequently during development. That property "
        "is part of why the project can credibly claim a maintainable engineering baseline "
        "even when predictive performance is not externally benchmarked."
    )

    # Table caption ABOVE the table (course rule)
    cap = doc.add_paragraph()
    cap.add_run(
        "Table 1. Unit test inventory by module (unittest, PYTHONPATH=src)."
    ).italic = True

    table = doc.add_table(rows=7, cols=4)
    table.style = "Table Grid"
    hdr = ("Module", "Topic", "Unit tests executed", "Outcome")
    for i, h in enumerate(hdr):
        table.rows[0].cells[i].text = h
    rows = [
        ("1", "First-order logic / matchup", "132", "Pass"),
        ("2", "Knowledge bases / defense", "67", "Pass"),
        ("3", "CSP assignment", "22", "Pass"),
        ("4", "Genetic batting order", "10", "Pass"),
        ("5", "Planning", "12", "Pass"),
        ("Total", "—", "243", "Pass"),
    ]
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            table.rows[r].cells[c].text = val

    doc.add_paragraph()
    doc.add_paragraph(
        "Integration tests (integration_tests.module3.test_module1_2_3_integration, "
        "integration_tests.module4.test_module3_4_integration, and "
        "integration_tests.module5.test_module1_2_3_4_5_integration) also passed, "
        "demonstrating that Module 3 consumes Modules 1–2 outputs, Module 4 consumes "
        "Module 3's roster, and Module 5 consumes the composed lineup objects without "
        "schema mismatch."
    )
    doc.add_paragraph(
        "Interpreting these results honestly, the strongest claim is internal consistency: "
        "the modules agree on data shapes and produce outputs that downstream components "
        "accept. A coach-facing evaluation would require separate study design (tasks, "
        "timers, error rates, subjective trust), which is outside the automated evidence "
        "collected here but is a natural extension if the project continues beyond the "
        "course."
    )
    doc.add_paragraph(
        "Figure 2. [Recommended] Screenshot of the Module 4–5 dashboard HTML after "
        "running the demo generator—use for visual evidence beyond automated tests. "
        "[Caption below figure.]"
    )

    doc.add_heading("Proposal Delta", level=1)
    doc.add_paragraph(
        "The approved proposal and the delivered system agree on the five-module topic "
        "stack and overall narrative (matchup logic, defensive knowledge base, CSP "
        "assignment, genetic order, planning). Documented shifts include: (1) Nine-slot "
        "roster model: proposal language sometimes emphasized eight fielders plus DH; "
        "implementation standardizes on P, C, 1B, 2B, 3B, SS, LF, CF, RF as the nine CSP "
        "variables, with pitcher defense treated as neutral in leverage scaling because "
        "Module 2 does not grade pitching defense. (2) Module 2 scope expansion: optional "
        "cross-position prediction improves sparse-data coverage at the cost of extra "
        "heuristic assumptions. (3) Module 5 inputs: informal game descriptions were refined "
        "into a typed game-state schema with explicit substitution accounting. "
        "(4) Demonstration surface: HTML dashboard and demo scripts support presentation "
        "without constituting a sixth AI module."
    )
    doc.add_paragraph(
        "One subtle delta concerns evaluation ambition: the proposal's narrative can read "
        "like end-to-end coaching automation, while the delivered artifact is closer to a "
        "pipeline library plus demo surfaces. That rescoping keeps claims testable and "
        "prevents the paper from overpromising outcomes that were never measured."
    )

    doc.add_heading("Limitations and Failure Analysis", level=1)
    doc.add_paragraph(
        "Limitation A — Heuristic scoring versus predictive validity: Modules 1–2 compress "
        "histories into bounded scores using rules. Small samples or missing fields can "
        "still yield confident-looking numbers unless operators add upstream data-quality "
        "checks. Improvement path: calibration against simulated run environments or "
        "labeled outcomes if the project expands beyond course scope."
    )
    doc.add_paragraph(
        "Limitation B — Optimization gaps: CSP assumes the nine-player input slice is "
        "correct; permissive eligibility can yield mathematically optimal but tactically "
        "odd assignments. The genetic algorithm may stagnate on fitness plateaus for flat "
        "rosters. Improvement path: adaptive mutation schedules and richer fitness terms "
        "with documented tuning methodology."
    )
    doc.add_paragraph(
        "Failure analysis for planning: when the bench is shallow or substitution limits "
        "are exhausted, the planner still returns structured output, but recommendations "
        "may converge to low-impact tactical guidance. That behavior is not necessarily a "
        "bug; it reflects constrained action spaces. The risk is user misinterpretation if "
        "confidence labels are read as endorsement rather than relative ranking within "
        "feasible actions."
    )

    doc.add_heading("Individual Contributions", level=1)
    doc.add_paragraph(
        "Replace the effort percentages below with the team-signed values required for "
        "grading (must total 100%). The 50/50 placeholder is not evidence-based."
    )
    doc.add_paragraph(
        "Kali Javetski — 50%: repository and checkpoint documentation; README and proposal-"
        "aligned specifications; cross-module integration and dashboard-related artifacts "
        "(verify against your team's internal log)."
    )
    doc.add_paragraph(
        "Sylvia Burroughs — 50%: parallel implementation streams reflected in MODULE3, "
        "MODULE4, and MODULE5 work-split documents; integration tests; planner and "
        "strategy rule separation (verify against your team's internal log)."
    )

    doc.add_heading("Conclusions and Future Work", level=1)
    doc.add_paragraph(
        "The team delivered a modular, test-backed lineup assistant that connects five AI "
        "topics into one pipeline with explicit interfaces and reproducible tests. The "
        "strongest empirical claim supported by this repository is software reliability "
        "under its own specifications, evidenced by 243 unit tests and three integration "
        "tests summarized in Table 1."
    )
    doc.add_paragraph(
        "Realistic next steps: (1) add an optional pitching or bullpen submodule, or remove "
        "unused fatigue fields to avoid misleading operators; (2) separate league-specific "
        "rules behind configuration; (3) run a coach-in-the-loop evaluation rather than "
        "only automated tests."
    )
    doc.add_paragraph(
        "If continued as a software product, the highest leverage improvement would likely "
        "be data integrity tooling (validation reports, confidence intervals, and clear "
        "missing-data behavior) before adding more optimization sophistication. Better "
        "inputs tend to improve every downstream stage without changing their algorithms."
    )

    doc.add_heading("References", level=1)
    refs = [
        "[1] Python Software Foundation, The Python Language Reference, Python 3.x "
        "documentation. [Online]. Available: https://docs.python.org/3/reference/",
        "[2] Python Software Foundation, unittest — Unit testing framework, Python 3.x "
        "documentation. [Online]. Available: https://docs.python.org/3/library/unittest.html",
        "[3] S. Slowinski, “Positional Adjustment,” FanGraphs Sabermetrics Library, FanGraphs. "
        "[Online]. Available: https://library.fangraphs.com/misc/war/positional-adjustment/",
    ]
    for line in refs:
        doc.add_paragraph(line, style="List Bullet")

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
